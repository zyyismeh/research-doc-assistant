"""多格式科研文档解析器 - 支持 PDF / Word / LaTeX / Markdown

技术要点：
- PyMuPDF + pdfplumber 实现 PDF 文本与表格提取
- python-docx 解析 Word 文档
- pylatexenc 解析 LaTeX 源文件
- 公式、图表、参考文献结构化提取
"""

from __future__ import annotations

import re
from pathlib import Path
from typing import Optional

from langchain_core.documents import Document
from loguru import logger


class DocumentParser:
    """统一文档解析入口，根据文件后缀自动选择解析策略"""

    SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".doc", ".tex", ".md", ".txt"}

    def parse(self, file_path: str | Path) -> list[Document]:
        file_path = Path(file_path)
        if not file_path.exists():
            raise FileNotFoundError(f"文件不存在: {file_path}")

        ext = file_path.suffix.lower()
        if ext not in self.SUPPORTED_EXTENSIONS:
            raise ValueError(f"不支持的文件格式: {ext}")

        logger.info(f"开始解析文档: {file_path.name} (格式: {ext})")

        parser_map = {
            ".pdf": self._parse_pdf,
            ".docx": self._parse_docx,
            ".doc": self._parse_docx,
            ".tex": self._parse_latex,
            ".md": self._parse_markdown,
            ".txt": self._parse_text,
        }
        docs = parser_map[ext](file_path)
        logger.info(f"文档解析完成: {file_path.name}, 提取 {len(docs)} 个文档片段")
        return docs

    def _parse_pdf(self, file_path: Path) -> list[Document]:
        """PDF 解析：PyMuPDF 提取文本 + pdfplumber 提取表格"""
        import fitz  # PyMuPDF

        documents: list[Document] = []
        pdf = fitz.open(str(file_path))

        for page_num in range(len(pdf)):
            page = pdf[page_num]
            text = page.get_text("text")

            if not text.strip():
                continue

            # 提取页面元数据
            metadata = {
                "source": str(file_path),
                "filename": file_path.name,
                "page": page_num + 1,
                "total_pages": len(pdf),
                "format": "pdf",
            }

            # 尝试提取标题（基于字体大小）
            blocks = page.get_text("dict")["blocks"]
            title = self._extract_title_from_blocks(blocks)
            if title:
                metadata["section_title"] = title

            documents.append(Document(page_content=text.strip(), metadata=metadata))

        pdf.close()

        # 用 pdfplumber 补充提取表格数据
        self._extract_tables_from_pdf(file_path, documents)

        return documents

    def _extract_title_from_blocks(self, blocks: list[dict]) -> Optional[str]:
        """从 PDF 文本块中基于字体大小提取疑似标题"""
        max_size = 0
        title = None
        for block in blocks:
            if block.get("type") != 0:  # 非文本块
                continue
            for line in block.get("lines", []):
                for span in line.get("spans", []):
                    if span["size"] > max_size and len(span["text"].strip()) < 200:
                        max_size = span["size"]
                        title = span["text"].strip()
        return title if max_size > 14 else None

    def _extract_tables_from_pdf(
        self, file_path: Path, documents: list[Document]
    ) -> None:
        """使用 pdfplumber 提取 PDF 表格并追加到文档列表"""
        try:
            import pdfplumber

            with pdfplumber.open(str(file_path)) as pdf:
                for page_num, page in enumerate(pdf.pages):
                    tables = page.extract_tables()
                    for table_idx, table in enumerate(tables):
                        if not table:
                            continue
                        table_text = self._table_to_markdown(table)
                        documents.append(
                            Document(
                                page_content=table_text,
                                metadata={
                                    "source": str(file_path),
                                    "filename": file_path.name,
                                    "page": page_num + 1,
                                    "content_type": "table",
                                    "table_index": table_idx,
                                    "format": "pdf",
                                },
                            )
                        )
        except Exception as e:
            logger.warning(f"pdfplumber 表格提取失败: {e}")

    @staticmethod
    def _table_to_markdown(table: list[list]) -> str:
        """将表格数据转换为 Markdown 格式"""
        if not table:
            return ""
        rows = []
        for i, row in enumerate(table):
            cells = [str(cell or "").strip() for cell in row]
            rows.append("| " + " | ".join(cells) + " |")
            if i == 0:
                rows.append("| " + " | ".join(["---"] * len(cells)) + " |")
        return "\n".join(rows)

    def _parse_docx(self, file_path: Path) -> list[Document]:
        """Word 文档解析"""
        from docx import Document as DocxDocument

        doc = DocxDocument(str(file_path))
        documents: list[Document] = []
        current_section = ""
        current_text: list[str] = []

        for para in doc.paragraphs:
            # 检测标题段落
            if para.style and para.style.name.startswith("Heading"):
                # 保存上一段内容
                if current_text:
                    documents.append(
                        Document(
                            page_content="\n".join(current_text),
                            metadata={
                                "source": str(file_path),
                                "filename": file_path.name,
                                "section_title": current_section,
                                "format": "docx",
                            },
                        )
                    )
                current_section = para.text.strip()
                current_text = []
            else:
                if para.text.strip():
                    current_text.append(para.text.strip())

        # 最后一段
        if current_text:
            documents.append(
                Document(
                    page_content="\n".join(current_text),
                    metadata={
                        "source": str(file_path),
                        "filename": file_path.name,
                        "section_title": current_section,
                        "format": "docx",
                    },
                )
            )

        return documents

    def _parse_latex(self, file_path: Path) -> list[Document]:
        """LaTeX 源文件解析 - 提取章节、公式、参考文献"""
        content = file_path.read_text(encoding="utf-8", errors="ignore")

        documents: list[Document] = []

        # 按 \section / \subsection 切分
        section_pattern = r"\\(?:section|subsection|subsubsection)\{([^}]+)\}"
        sections = re.split(section_pattern, content)

        current_title = "前言"
        for i, part in enumerate(sections):
            if i % 2 == 1:  # 标题
                current_title = part.strip()
            else:
                clean_text = self._clean_latex(part)
                if clean_text.strip():
                    documents.append(
                        Document(
                            page_content=clean_text.strip(),
                            metadata={
                                "source": str(file_path),
                                "filename": file_path.name,
                                "section_title": current_title,
                                "format": "latex",
                            },
                        )
                    )

        # 提取公式
        formulas = re.findall(
            r"\\begin\{(?:equation|align)\*?\}(.*?)\\end\{(?:equation|align)\*?\}",
            content,
            re.DOTALL,
        )
        for idx, formula in enumerate(formulas):
            documents.append(
                Document(
                    page_content=f"数学公式: {formula.strip()}",
                    metadata={
                        "source": str(file_path),
                        "filename": file_path.name,
                        "content_type": "formula",
                        "formula_index": idx,
                        "format": "latex",
                    },
                )
            )

        return documents

    @staticmethod
    def _clean_latex(text: str) -> str:
        """清理 LaTeX 命令，保留纯文本"""
        text = re.sub(r"\\(?:begin|end)\{[^}]+\}", "", text)
        text = re.sub(r"\\[a-zA-Z]+\{([^}]*)\}", r"\1", text)
        text = re.sub(r"\\[a-zA-Z]+", "", text)
        text = re.sub(r"[{}]", "", text)
        text = re.sub(r"%.*$", "", text, flags=re.MULTILINE)
        text = re.sub(r"\n{3,}", "\n\n", text)
        return text

    def _parse_markdown(self, file_path: Path) -> list[Document]:
        """Markdown 文件解析 - 按标题切分章节"""
        content = file_path.read_text(encoding="utf-8")
        documents: list[Document] = []

        sections = re.split(r"(^#{1,3}\s+.+$)", content, flags=re.MULTILINE)

        current_title = ""
        for i, part in enumerate(sections):
            if re.match(r"^#{1,3}\s+", part):
                current_title = part.strip("# \n")
            elif part.strip():
                documents.append(
                    Document(
                        page_content=part.strip(),
                        metadata={
                            "source": str(file_path),
                            "filename": file_path.name,
                            "section_title": current_title,
                            "format": "markdown",
                        },
                    )
                )

        return documents

    def _parse_text(self, file_path: Path) -> list[Document]:
        """纯文本解析"""
        content = file_path.read_text(encoding="utf-8", errors="ignore")
        return [
            Document(
                page_content=content,
                metadata={
                    "source": str(file_path),
                    "filename": file_path.name,
                    "format": "text",
                },
            )
        ]


class AcademicTextCleaner:
    """学术文本清洗工具 - 去噪、分词、元数据抽取"""

    @staticmethod
    def clean(text: str) -> str:
        """学术文本清洗"""
        # 去除页眉页脚噪声
        text = re.sub(r"第\s*\d+\s*页.*?共\s*\d+\s*页", "", text)
        text = re.sub(r"Page\s+\d+\s+of\s+\d+", "", text, flags=re.IGNORECASE)
        # 归一化空白字符
        text = re.sub(r"[ \t]+", " ", text)
        text = re.sub(r"\n{3,}", "\n\n", text)
        # 去除参考文献编号噪声
        text = re.sub(r"\[\d+(?:,\s*\d+)*\]", "", text)
        return text.strip()

    @staticmethod
    def extract_references(text: str) -> list[str]:
        """提取参考文献列表"""
        ref_pattern = r"\[(\d+)\]\s*(.+?)(?=\[\d+\]|\Z)"
        return [match[1].strip() for match in re.findall(ref_pattern, text, re.DOTALL)]

    @staticmethod
    def segment_chinese(text: str) -> str:
        """中文分词（用于 BM25 索引）"""
        import jieba

        return " ".join(jieba.cut(text))
